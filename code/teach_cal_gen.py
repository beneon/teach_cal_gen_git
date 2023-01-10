
from docx import Document
import pandas as pd
from yaml import safe_load
from datetime import date, datetime, timedelta
import os
import re
import snoop


entry_index:int = 1

with open('config.yml',encoding='utf8') as _yaml_file:
    _config = safe_load(_yaml_file)
semester_date_str = _config['semester_date_str']
weeks_yaml_file_path = _config['weeks_yaml_file_name']
weeks_yaml_file_path = os.path.join('..', 'datafile', semester_date_str, weeks_yaml_file_path)
doc_template_path = _config['teach_cal_template_docx']
doc_template_path = os.path.join('..','datafile',semester_date_str,doc_template_path)
if entry_index < len(_config['teach_cal_data_list']):
    course_data_entry = _config['teach_cal_data_list'][entry_index]['entry']
else:
    raise Exception('entry_index > teach_cal_data_list length')
xl_data_path = os.path.join('..','output',semester_date_str,_config['teach_cal_time_table_xlsx'])
course_yaml_data_path = os.path.join('..','datafile',semester_date_str,course_data_entry['course_data'])
course_name = course_data_entry['course_name']
doc = Document(doc_template_path)
xlFile = pd.ExcelFile(xl_data_path)
dtypes = {
    '节数':str,
    '节次':str,
    '周数':int,
    '星期几':int,
}


doc_shadow = None

class CourseCalGenerator:

    weekday_int2str = ['星期日','星期一','星期二','星期三','星期四','星期五','星期六']
    digit2character = ['零','一','二','三','四','五','六','七','八','九','十','十一','十二','十三','十四','十五','十六','十七','十八','十九','二十']
    type_desc = {
        'lecture':'理论课',
        'discussion':'讨论课',
        'experiment':'实验',
        'exercise':'习题课',
        'other':'其他',
    }
    course_type_list = list(type_desc.keys())
    reDigit = re.compile(r'\d{2}')

    with open(weeks_yaml_file_path, encoding='utf8') as data_file:
        weeks_data = safe_load(data_file)

    def __init__(self, timetable:pd.DataFrame, course_data:dict):
        self.doc = Document(doc_template_path)
        self.timetable = timetable
        self.timetable['星期几'] = self.timetable['星期几'].apply(lambda e:self.weekday_int2str[int(e)])
        self.timetable['周数说明'] = self.timetable['周数'].apply(lambda e:self.week_desc_gen(e))
        self.timetable['节数'] = self.timetable['节次'].apply(lambda e:len(e)/2)
        self.timetable['节数'] = self.timetable['节数'].cumsum()
        self.course_data = course_data
        # 处理course_data, 提取部分不方便直接访问的数据
        teachers_data:list = self.course_data['teachers']
        self.teachers_name = [e['teacher']['name'] for e in teachers_data]
        self.teachers_name = self.join_list_considering_len1_list(self.teachers_name)
        self.teachers_title = [e['teacher']['title'] for e in teachers_data]
        self.teachers_title = self.join_list_considering_len1_list(self.teachers_title)
        self.df_course_sessions = pd.DataFrame([e['session'] for e in self.course_data['course_detail']])
        # 生成记录, 当前session是对应类型的第几次(1 based int)
        grp = self.df_course_sessions.groupby('type')
        self.df_course_sessions['ith_session'] = grp.cumcount()+1
        # 处理timetable, 生成若干统计数据
        self.course_weeks = len(self.timetable['周数'].unique())

        # # 处理文档模板:
        self.doc.paragraphs[2].text = self.weeks_data['semester_name']
        self.doc = self.doc_paragraph_process(self.doc)
        self.doc = self.doc_table_reference_process(self.doc)
        self.doc = self.doc_table_time_summary(self.doc)
        self.doc = self.doc_table_course_schedule(self.doc)

    def doc_paragraph_process(self,doc:Document) -> Document:
        """
        process the upper left section, including:
        should be in paragraph 4
        :param doc:Document doc object in the class
        :return: processed doc object
        """
        assert "授课教师姓名" in doc.paragraphs[4].text, f"模板格式有改动, 基本信息不在第五段:{doc.paragraphs[4].text}"
        # 数据未到, 暂时空置
        para1_classes = self.timetable.loc[0,'班级']
        doc.paragraphs[4].text = f"""授课教师姓名  {self.teachers_name}
职称  {self.teachers_title}
授课专业班级  {para1_classes}
课程名称    {self.course_data['title']}
根据何种教学大纲   {self.course_data['curriculum_type']}
采用教材名称     {self.course_data['text_book']}
"""

        return doc

    def doc_table_reference_process(self,doc:Document) -> Document:
        """
        tables[2] should be the table, with 1 row, 2 column
        '参考书目的格式：\n{tbl3.references}\n
        :param doc:Document doc object in the class
        :return:Document processed doc object
        """
        assert len(doc.tables)==3, '文档中表格数量不为3'
        tbl = doc.tables[2]
        assert len(tbl.row_cells(0))==2, '第三个表格列数不为2'
        ref_txt = "参考书目\n"
        ref = self.course_data['references']
        ref = [f"[{i+1}] {e}\n" for i,e in enumerate(ref)]
        ref_txt += "".join(ref)
        tbl.row_cells(0)[1].text = ref_txt
        doc.tables[2] = tbl
        return doc

    def time_summary(self,df_sessions,nafiller=0):
        df_groupby = df_sessions.groupby('type')
        df_summary = df_groupby.sum()
        df_summary_tobe_filled = pd.DataFrame({'type': self.course_type_list}).set_index('type')
        df_summary = df_summary_tobe_filled.join(df_summary).fillna(nafiller)
        return df_summary


    def doc_table_time_summary(self,doc:Document):
        """
        右上部分的文本替换
        :param doc:
        :return:
        """
        assert len(doc.tables[0].rows) == 1 & len(doc.tables[0].row_cells(0)) == 1, '第一个表格格式不正确'
        rc = doc.tables[0].row_cells(0)[0]

        df_summary = self.time_summary(self.df_course_sessions)
        rc_txt = f"""
周  数  {self.course_weeks}  周
讲课{df_summary.loc['lecture','class_span']:.0f} 学时  课堂讨论  {df_summary.loc['discussion','class_span']:.0f}  学时
训练课 {df_summary.loc['experiment','class_span']:.0f}   学时  习 题 课  {df_summary.loc['exercise','class_span']:.0f}  学时 
其他环节  {df_summary.loc['other','class_span']:.0f}    学时
总 计  {df_summary['class_span'].sum():.0f}  学时
"""
        rc.text = rc_txt
        doc.tables[0].row_cells(0)[0] = rc
        return doc

    def doc_table_course_schedule(self, doc:Document) -> Document:
        """
        根据
        :param doc: Document
        :return: Document
        """

        # 这里有一个逻辑要写进去:
        # 1. 首先, 你要确定这次课他有几节, 节数和节次都是关于节数的信息, 节次是最好的, 但是节数也是不错, 不过说真的, class_span这种东西, 每次课应该是一样的, 真要是有什么特殊情况, 也应该是可以手动修改的. 那么你要怎么做呢? 我觉得最好是在excel表格中修正一下节次, 如果节次为空, 那就手动填上去. 根据节次信息来计算当前课程单元有几节课
        # 2. 节数确定以后, 一周一周遍历, 每一周里面, 一次课程一次课程遍历, 每次课程都是一个slot, 这个slot有固定的容量(class_span), 然后从教学日历yaml数据文件里面(self.df_course_sessions)去遍历那边的课程安排, 一个单元的课程一定要在一次课里安排完, 但是一次课可以安排多个单元的课程
        # 3. 安排完成以后,接下来才是将一周所有的课程安排组装成描述文本. 之后通过docx的api新建一行,填充各个单元格

        # 实际流程: 首先是遍历df_course_sessions, 计算一个节数的cumsum, 另一方面遍历timetable, 也产生一个节数的cumsum, 将两者相比, session的cumsum往timetable的cumsum中插队
        tbl = doc.tables[1]
        df_course_sessions = self.df_course_sessions
        df_course_sessions['class_span_sum'] = df_course_sessions['class_span'].cumsum()
        course_span_sum_ttl = df_course_sessions['class_span_sum'].to_list()[-1]
        timetable_span_ttl = self.timetable['节数'].to_list()[-1]
        assert course_span_sum_ttl==timetable_span_ttl, f"total course span for timetable({timetable_span_ttl}) and course data ({course_span_sum_ttl}) does not match"
        df_course_sessions['fk'] = df_course_sessions['class_span_sum'].apply(self.cumsum_compare)
        print(df_course_sessions[['fk','class_span_sum']],self.timetable.index)
        df_joined = df_course_sessions.set_index('fk').join(self.timetable,how='outer',lsuffix='course',rsuffix='time')
        tt_grp_by = df_joined.groupby('周数')
        course_weeks = list(tt_grp_by.indices)
        weeks = range(1,20)
        # 有人说，不用保留空行，但是谁知道呢，万一哪天哪个神经病说空行还是需要的呢？下面一行是应对这种情况的
        # for week in weeks:
        # 所以下面两行代码看起来怪怪的，就是为了应对在weeks中遍历的情况
        for week in course_weeks:
            if week in course_weeks:
                df_thisweek = tt_grp_by.get_group(week)
                df_thisweek['desc'] = df_thisweek.apply(lambda r:self.gen_course_schedule_desc(r),axis=1)
                week_desc = df_thisweek['周数说明'].values[0]
                srs_week_time_summary = self.time_summary(df_thisweek, nafiller=0)['class_span']
                _row_added = tbl.add_row()
                assert len(_row_added.cells)==8, '主表格式错误, 没有8列'
                _row_added.cells[0].text = week_desc
                _row_added.cells[1].text = f"{srs_week_time_summary.sum():.0f}"

                def zero2blank(e):
                    if e == 0:
                        return ""
                    else:
                        return f"{e:.0f}"

                srs_week_time_summary = srs_week_time_summary.apply(zero2blank)
                _row_added.cells[2].text = srs_week_time_summary['lecture']
                _row_added.cells[3].text = srs_week_time_summary['experiment']
                _row_added.cells[4].text = srs_week_time_summary['exercise']
                _row_added.cells[5].text = srs_week_time_summary['discussion']
                _row_added.cells[6].text = srs_week_time_summary['other']
                # snoop.pp(df_thisweek['desc'])
                _row_added.cells[7].text = df_thisweek['desc'].str.cat(sep='\n')
            else:
                week_desc = self.week_desc_gen(week)
                _row_added = tbl.add_row()
                _row_added.cells[0].text = week_desc



        return doc

    def join_list_considering_len1_list(self, list_src:list):
        if len(list_src) == 1:
            return list_src[0]
        else:
            return ', '.join(list_src)

    def week_desc_gen(self,week_num):
        day1 = self.weeks_data['1st_day']
        day1_this_week = day1 + timedelta(days=(week_num-1)*7)
        day7_this_week = day1_this_week + timedelta(days=6)
        memo = self.weeks_data['weeks'][week_num-1]['week']['memo']
        week_num_chr = f"""
第{self.digit2character[week_num]}周
{day1_this_week:%m}月{day1_this_week:%d}日
至
{day7_this_week:%m}月{day7_this_week:%d}日
"""
        if memo:
            week_num_chr += memo
        return week_num_chr

    def cumsum_compare(self,course_cumsum):
        if course_cumsum <= self.timetable.loc[0,'节数']:
            return 0
        else:
            comparison_srs = self.timetable['节数'].ge(course_cumsum)
            comparison_srs_shift1 = comparison_srs.shift(1).fillna(False)
            comparison_multiplied =comparison_srs != comparison_srs_shift1
            return comparison_srs[comparison_multiplied].index[0]

    def gen_course_schedule_desc(self,course_row):
        # 将节次这种0405的文本转换成4,5这样的数值
        # 之前这种写法……明显有错啊，肯定要用re findall啊
        # jieci_list = course_row['节次'].split('0')[1:]
        jieci_list = self.reDigit.findall(course_row['节次'])
        assert len(jieci_list)>0, f"{course_row}的节次信息无法提取数据"
        jieci_list = [int(e) for e in jieci_list]
        start_jie = min(jieci_list)
        end_jie = max(jieci_list)
        if end_jie-start_jie==1:
            jieci_sep = ','
        else:
            jieci_sep = '-'

        # 根据session的content生成content文本
        content = ""
        if type(course_row['content'])==str:
            content = course_row['content']
        elif type(course_row['content'])==list:
            for i,e in enumerate(course_row['content']):
                content += f"{e}\n"
        else:
            content=""


        print(course_row)
        rst_txt = f"""
讲授教师：{course_row['teacher']}({course_row['class_span']}学时)  {course_row['星期几']}，{start_jie}{jieci_sep}{end_jie}节
课程介绍
{self.type_desc[course_row['type']]}{self.digit2character[course_row['ith_session']]} {course_row['title']}
{content}
"""
        return rst_txt

# with open('spss_experiment.yaml') as data_file:
#     spss_course_data = full_load(data_file)

# cgen = CourseCalGenerator(df_kouchuang_spss,spss_course_data)
# cgen.doc.save('17 口腔.docx')


df = xlFile.parse(sheet_name=entry_index,dtype=dtypes)
with open(course_yaml_data_path,encoding='utf8') as data_dict_file:
    data_dict = safe_load(data_dict_file)
cgen = CourseCalGenerator(df,data_dict)
cgen.doc.save(os.path.join('..','output',semester_date_str,f'{course_name}.docx'))